from pathlib import Path
import argparse
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BODY_FONT = "Songti SC"
HEADING_FONT = "Heiti SC"
NUMERIC_CITATION_RE = re.compile(
    r"\[(?:\d+(?:-\d+)?)(?:\s*,\s*\d+(?:-\d+)?)*\]"
)
CITATION_STYLES = (
    "superscript-number",
    "superscript-bracket",
    "inline-bracket",
    "inline-parenthesis",
    "author-date",
)
REFERENCE_LABEL_STYLES = ("dot", "bracket", "none")


def set_font(run, font=BODY_FONT, size=10.5, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_para(paragraph, size=10.5, font=BODY_FONT, bold=False, align=None, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_font(run, font, size, bold)


def set_section(section, landscape=False):
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def set_cell_border(cell, edge, val="single", size="4", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    tag = "w:" + edge
    border = borders.find(qn(tag))
    if border is None:
        border = OxmlElement(tag)
        borders.append(border)
    border.set(qn("w:val"), val)
    if val != "nil":
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "nil")
    for edge, size in (("top", "10"), ("bottom", "10")):
        border = borders.find(qn("w:" + edge))
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_cell_padding(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_markup_runs(paragraph, text, size=10.5, bold=False, citation_style=None):
    pattern = re.compile(
        r"(\*[^*]+\*|\[(?:\d+(?:-\d+)?)(?:\s*,\s*\d+(?:-\d+)?)*\])"
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Times New Roman", size, bold)
            run.italic = True
        elif NUMERIC_CITATION_RE.fullmatch(part):
            if citation_style in {None, "author-date"}:
                raise ValueError(
                    "Numeric [n] citation found without a numeric journal citation style"
                )
            number = part[1:-1]
            labels = {
                "superscript-number": number,
                "superscript-bracket": part,
                "inline-bracket": part,
                "inline-parenthesis": f"({number})",
            }
            superscript = citation_style.startswith("superscript-")
            run = paragraph.add_run(labels[citation_style])
            run_size = max(7.0, size * 0.78) if superscript else size
            set_font(run, "Times New Roman", run_size, False)
            run.font.superscript = superscript
        else:
            for token in re.split(r"([\x20-\x7e]+)", part):
                if not token:
                    continue
                run = paragraph.add_run(token)
                token_font = "Times New Roman" if re.fullmatch(r"[\x20-\x7e]+", token) else BODY_FONT
                set_font(run, token_font, size, bold)


def set_cell_text(cell, text, size=8.5, bold=False, citation_style=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = None
    add_markup_runs(p, text, size, bold, citation_style)
    set_cell_padding(cell)


def markdown_table(rows, doc, wide=False, citation_style=None):
    header = rows[0]
    body = rows[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable_cm = 25.0 if wide else 15.8
    widths = column_widths(header, usable_cm)
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.width = Cm(widths[i])
        set_cell_text(
            cell,
            text,
            size=8 if wide else 8.5,
            bold=True,
            citation_style=citation_style,
        )
        set_cell_border(cell, "bottom", size="5")
    set_repeat_table_header(table.rows[0])
    set_cant_split(table.rows[0])
    for row in body:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Cm(widths[i])
            set_cell_text(
                cells[i],
                text,
                size=8.0 if wide else 8.6,
                citation_style=citation_style,
            )
        set_cant_split(table.rows[-1])
    set_table_borders(table)
    set_table_widths(table, widths)
    return table


def set_table_widths(table, widths_cm):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(Cm(w).twips) for w in widths_cm)))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(width).twips)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(Cm(width).twips)))


def column_widths(header, usable_cm):
    n = len(header)
    if n == 6:
        return [4.8, 2.4, 4.2, 3.2, 7.0, 3.4]
    if n == 5:
        return [2.4, 3.4, 4.2, 2.6, 3.2]
    if n == 4:
        return [2.2, 5.0, 3.6, 5.0]
    return [usable_cm / n] * n


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        line = lines[i].strip()
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, HEADING_FONT, 13 if level == 1 else 11.5, True)
    return p


def add_paragraph_with_italics(
    doc, text, body_size=10.5, citation_style=None
):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_markup_runs(p, text, body_size, False, citation_style)
    return p


def add_reference_paragraph(
    doc, line, reference_label_style, citation_style
):
    match = re.match(r"^\[(\d+)\]\s*(.*)$", line)
    if not match:
        return add_paragraph_with_italics(doc, line, 9, citation_style)
    if reference_label_style == "none":
        raise ValueError(
            "Numbered [n] reference found while reference label style is none"
        )
    number, text = match.groups()
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.65)
    labels = {"dot": f"{number}. ", "bracket": f"[{number}] "}
    run = p.add_run(labels[reference_label_style])
    set_font(run, "Times New Roman", 9, False)
    add_markup_runs(p, text, 9, False, citation_style)
    return p


def validate_citation_configuration(lines, citation_style, reference_label_style):
    text = "\n".join(lines)
    has_numeric_markers = bool(NUMERIC_CITATION_RE.search(text))
    if citation_style == "author-date":
        if reference_label_style != "none":
            raise ValueError(
                "Author-date citations require --reference-label-style none"
            )
        if has_numeric_markers:
            raise ValueError(
                "Author-date manuscript still contains canonical numeric [n] markers"
            )
    elif reference_label_style == "none":
        raise ValueError(
            "Numeric citation styles require dot or bracket reference labels"
        )


def build(src, out, citation_style, reference_label_style):
    doc = Document()
    set_section(doc.sections[0], landscape=False)
    styles = doc.styles
    styles["Normal"].font.name = BODY_FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    styles["Normal"].font.size = Pt(10.5)

    lines = src.read_text(encoding="utf-8").splitlines()
    validate_citation_configuration(lines, citation_style, reference_label_style)
    table_count = 0
    in_refs = False
    landscape_active = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line[2:])
            set_font(run, HEADING_FONT, 18, True)
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:]
            if heading == "参考文献":
                in_refs = True
            if landscape_active and not (heading.startswith("2 ") or heading.startswith("参考文献")):
                set_section(doc.add_section(WD_SECTION.NEW_PAGE), landscape=False)
                landscape_active = False
            add_heading(doc, heading, 1)
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            table_count += 1
            wide = table_count == 2
            if wide and not landscape_active:
                set_section(doc.add_section(WD_SECTION.NEW_PAGE), landscape=True)
                landscape_active = True
            if not wide and landscape_active:
                set_section(doc.add_section(WD_SECTION.NEW_PAGE), landscape=False)
                landscape_active = False
            markdown_table(
                rows,
                doc,
                wide=wide,
                citation_style=citation_style,
            )
            if wide:
                doc.add_paragraph()
            continue

        if re.match(r"^表2", line) and not landscape_active:
            set_section(doc.add_section(WD_SECTION.NEW_PAGE), landscape=True)
            landscape_active = True

        p = (
            add_reference_paragraph(
                doc, line, reference_label_style, citation_style
            )
            if in_refs
            else add_paragraph_with_italics(
                doc, line, 10.5, citation_style
            )
        )
        if line.startswith("作者") or line.startswith("单位"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
        if re.match(r"^表\d", line):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.bold = True
        i += 1

    doc.save(out)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build a Chinese bencao review DOCX from Markdown.")
    parser.add_argument("input", type=Path, help="Markdown manuscript path")
    parser.add_argument("output", type=Path, nargs="?", help="Output DOCX path")
    parser.add_argument(
        "--citation-style",
        required=True,
        choices=CITATION_STYLES,
        help="Inline citation rendering from the verified journal profile",
    )
    parser.add_argument(
        "--reference-label-style",
        required=True,
        choices=REFERENCE_LABEL_STYLES,
        help="Reference-list label style from the verified journal profile",
    )
    args = parser.parse_args()
    output = args.output or args.input.with_name(args.input.stem + "_投稿稿.docx")
    build(
        args.input,
        output,
        args.citation_style,
        args.reference_label_style,
    )
    print(output)
